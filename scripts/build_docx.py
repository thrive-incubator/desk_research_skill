#!/usr/bin/env python3
"""Render a desk research Markdown report into a Thrive-branded .docx.

Usage:
    python3 build_docx.py --input report.md --output report.docx

The script is intentionally forgiving: it parses standard Markdown produced by
the thrive-desk-research skill (see references/report-template.md) and renders:
  - `# H1`  -> document title (teal, large)
  - `## H2` -> section heading (teal)
  - `### H3` -> subsection heading
  - `| ... |` Markdown tables -> Word tables with a teal header row
  - `- ` / `* ` bullets, `1.` numbered lists
  - `**bold**` inline emphasis
  - everything else -> body paragraph

It does NOT try to reproduce the original template's multi-column boxed layout
(Word doesn't do that cleanly from a script); it produces a clean, on-brand
linear document that carries the same content and section order. Branding is a
single teal accent so it reads as a Thrive artifact, not a generic export.
"""

import argparse
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")

THRIVE_TEAL = RGBColor(0x1F, 0x7A, 0x7A)  # matches the template's teal headings
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(paragraph, text):
    """Add text to a paragraph, honoring **bold** and [link](url) markup."""
    # Flatten links to "text (url)" so citations survive in the docx.
    text = LINK.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    pos = 0
    for m in INLINE_BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def style_heading(paragraph, size, color=THRIVE_TEAL, bold=True):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold


def parse_table(lines, start):
    """Parse a contiguous block of Markdown table rows starting at `start`."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # Skip the |---|---| separator row.
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def render(md_text, doc):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        if c < len(table.rows[r].cells):
                            cell = table.rows[r].cells[c]
                            cell.text = ""
                            add_inline(cell.paragraphs[0], cell_text)
                            if r == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                                    run.font.color.rgb = THRIVE_TEAL
                doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:])
            style_heading(p, 22)
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:])
            style_heading(p, 15)
        elif stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:])
            style_heading(p, 12)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("_") and stripped.endswith("_") and len(stripped) > 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("_"))
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)

        i += 1


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="Markdown report path")
    ap.add_argument("--output", required=True, help="Output .docx path")
    args = ap.parse_args()

    with open(args.input, encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    render(md_text, doc)
    doc.save(args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
